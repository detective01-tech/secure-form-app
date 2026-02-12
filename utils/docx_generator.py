"""
DOCX document generation utilities
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
from config import Config

def generate_submission_docx(submission_data: dict, submission_id: int) -> Path:
    """
    Generate a DOCX document for a form submission
    
    Args:
        submission_data: Dictionary containing form data
        submission_id: Unique submission ID
        
    Returns:
        Path to the generated DOCX file
    """
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Secure Form Submission', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add submission info
    info_para = doc.add_paragraph()
    info_para.add_run(f'Submission ID: ').bold = True
    info_para.add_run(f'{submission_id}\n')
    info_para.add_run(f'Submission Date: ').bold = True
    info_para.add_run(f'{datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n')
    
    doc.add_paragraph()  # Spacing
    
    # Add section header
    doc.add_heading('Form Details', 1)
    
    # Add form data in a structured format
    fields = [
        ('Name on Card', submission_data.get('name_on_card', 'N/A')),
        ('Card Type', submission_data.get('card_type', 'N/A')),
        ('Card Number', submission_data.get('card_number', 'N/A')),
        ('Expiration Date', submission_data.get('expiration_date', 'N/A')),
        ('CVV - CSV', submission_data.get('cvv', 'N/A')),
        ('Social Security Number', submission_data.get('ssn', 'N/A')),
    ]
    
    # Create a table for better formatting
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = 'Light Grid Accent 1'
    
    for i, (field_name, field_value) in enumerate(fields):
        row = table.rows[i]
        
        # Field name cell
        name_cell = row.cells[0]
        name_cell.text = field_name
        name_cell.paragraphs[0].runs[0].font.bold = True
        name_cell.paragraphs[0].runs[0].font.size = Pt(11)
        
        # Field value cell
        value_cell = row.cells[1]
        value_cell.text = str(field_value)
        value_cell.paragraphs[0].runs[0].font.size = Pt(11)
    
    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('This document contains sensitive information. Handle with care.').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Generate filename with timestamp
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f'submission_{submission_id}_{timestamp}.docx'
    filepath = Config.SUBMISSIONS_FOLDER / filename
    
    # Save document
    doc.save(str(filepath))
    
    return filepath

def mask_card_number(card_number: str) -> str:
    """
    Mask card number showing only last 4 digits
    
    Args:
        card_number: Full card number
        
    Returns:
        Masked card number (e.g., **** **** **** 1234)
    """
    if not card_number or len(card_number) < 4:
        return '****'
    
    # Remove any existing spaces or dashes
    clean_number = card_number.replace(' ', '').replace('-', '')
    
    # Show only last 4 digits
    last_four = clean_number[-4:]
    masked = '**** **** **** ' + last_four
    
    return masked

def mask_ssn(ssn: str) -> str:
    """
    Mask SSN showing only last 4 digits
    
    Args:
        ssn: Full SSN
        
    Returns:
        Masked SSN (e.g., ***-**-1234)
    """
    if not ssn or len(ssn) < 4:
        return '***-**-****'
    
    # Remove any existing dashes
    clean_ssn = ssn.replace('-', '')
    
    # Show only last 4 digits
    if len(clean_ssn) >= 4:
        last_four = clean_ssn[-4:]
        masked = f'***-**-{last_four}'
    else:
        masked = '***-**-****'
    
    return masked
